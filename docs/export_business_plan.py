#!/usr/bin/env python3
"""Export 商业策划书 to Word (.docx) and PDF with embedded images."""
from __future__ import annotations

import os
import re
import shutil
import subprocess
import sys
from pathlib import Path
from urllib.parse import urlparse

import pypandoc
import requests

DOCS = Path(__file__).resolve().parent
SRC_MD = DOCS / "商业策划书_HDW建材城.md"
ASSETS = DOCS / "assets"
OUT_DOCX = DOCS / "HDW商业策划书.docx"
OUT_PDF = DOCS / "HDW商业策划书.pdf"
WORK_MD = DOCS / "_export_work.md"

IMAGE_URLS = [
    ("hero", "https://img.thebdsoft.com/202510/c4a282869899a386a3c5fc43912c82d0.jpg", "HDW 工程案例 Hero"),
    ("logo", "https://img.thebdsoft.com/202505/5eb1be0a9abf125781fd3f2683bbbbd3.jpg", "HDW 品牌 Logo"),
    ("about", "https://img.thebdsoft.com/202505/ae9af419394511098c4d1ea9c0a4e2c5.jpg", "HDW 工程案例"),
    ("design", "https://img.thebdsoft.com/202507/7ac03599c211225438ece21d893508f9.png", "专业设计方案"),
    ("quote", "https://img.thebdsoft.com/202507/a44328600b3ce6ed5dd1572139d14650.png", "透明报价流程"),
    ("pm", "https://img.thebdsoft.com/202510/f5023528dd15e51b7f00b58006db139b.jpg", "高效项目管理"),
    ("qc", "https://img.thebdsoft.com/202510/1184004662612f64d8b25ac4e90367e1.jpg", "严格品质控制"),
    ("service", "https://img.thebdsoft.com/202510/fd60e34815ec1022e2fb81d72026b59c.jpg", "贴心客户服务"),
    ("floor", "https://www.hdwbuild.com/uploads/20260331/img_1774946807180_8013_lqp1uvttjy9.jpg", "地板 SKU"),
    ("tile", "https://www.hdwbuild.com/uploads/20260331/img_1774946921533_7093_ftz7fe7jl8.jpg", "瓷砖 SKU"),
]


def log(msg: str) -> None:
    sys.stdout.buffer.write((msg + "\n").encode("utf-8", errors="replace"))


def download_images() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    for name, url, _ in IMAGE_URLS:
        ext = Path(urlparse(url).path).suffix or ".jpg"
        dest = ASSETS / f"{name}{ext}"
        if not dest.exists():
            log(f"Downloading {name}...")
            r = requests.get(url, timeout=60, headers={"User-Agent": "HDW-Export/1.0"})
            r.raise_for_status()
            dest.write_bytes(r.content)
        paths[name] = dest
    return paths


def rel(p: Path) -> str:
    return p.relative_to(DOCS).as_posix()


def remove_milestone_section(text: str) -> str:
    """Remove §11.4 for Word/PDF export only (source MD unchanged)."""
    return re.sub(
        r"### 11\.4 里程碑对赌[^\n]*\n\n\| 里程碑.*?\n\n---\n\n",
        "",
        text,
        count=1,
        flags=re.DOTALL,
    )


def mermaid_to_prose(snippet: str) -> str:
    """Convert mermaid blocks to readable Chinese prose for Word/PDF."""
    s = snippet.strip()
    if "imagineBoard" in s:
        return (
            "用户在建材卡片点击 ✨，将最多 6 种材料加入 imagineBoard。"
            "在 ImaginePanel 选择房间类型与设计风格后，前端调用 POST /api/imagine；"
            "后端使用 OpenAI gpt-image-1 生成效果图并返回给用户预览。"
        )
    if "Phase 0" in s and "LA/OC/SGV" in s:
        return (
            "扩张按四个阶段推进：**Phase 0（0–12 月）** 聚焦南加州 LA、OC、SGV；"
            "**Phase 1（12–24 月）** 进入北加州湾区、圣地亚哥与拉斯维加斯华人圈；"
            "**Phase 2（24–36 月）** 拓展西雅图、休斯顿/达拉斯、纽约/法拉盛及温哥华/多伦多；"
            "**Phase 3（36 月以上）** 覆盖全美 Top 20 亚裔都市圈，并探索国内高净值跨境定制。"
            "各阶段以南加州验证模型为起点，逐步向东西海岸及加拿大华人社区复制。"
        )
    if "小红书/Google/口碑" in s:
        return (
            "用户通过小红书、Google 搜索或口碑推荐进入 HDW 商城，浏览并搜索 SKU；"
            "可将建材加入 AI 样板间生成效果图。若满意则加入购物车并下单支付。"
            "履约分两种：本地现货由 LA 仓发货或自提，国内直采则经海运、清关后配送。"
            "最终由 HDW 推荐安装或自有施工完成交付，用户评价后形成复购与转介绍。"
        )
    if "sequenceDiagram" in s:
        return (
            "承包商向 HDW 平台上传项目清单或 BOM，平台在 24 小时内返回中文报价单。"
            "承包商确认 PO 并支付定金后：若本地有现货，仓库/3PL 拣货并配送至工地；"
            "若需国内订货，HDW 向国内供应商下单排产，海运至 LA 仓后再分批配送。"
            "项目结束后 HDW 开具发票并提供售后支持。"
        )
    if 'Client["浏览器' in s or ("ImaginePanel" in s and "AdminPanel" in s):
        return (
            "前端为 React 商城界面，包含 ImaginePanel（AI 样板间）与 AdminPanel（商城管理）。"
            "Express API 提供商品接口、AI 生成接口、管理员接口及本地上传图片静态服务。"
            "OpenAI gpt-image-1 负责图像生成，hdwbuild.com CDN 提供品牌与产品素材；"
            "商品数据存于 products.json，运营上传图片存于 server/uploads。"
        )
    if 'CN["中国产业带"]' in s or "佛山瓷砖" in s:
        return (
            "中国产业带（佛山瓷砖、湖州地板、南安卫浴、中山灯具）经货代海运至 LA，"
            "由清关 Broker 处理后入 LA 中央仓。仓内库存分别支撑 HDW 商城 D2C 零售"
            "与 B2B 项目配送；部分大件 OEM 产品可跳过中央仓、直发客户（Dropship 少数 SKU）。"
        )
    if "SKU 类型" in s:
        return (
            "订单按 SKU 类型分流：**A 类**高周转标准品走 LA 仓现货（24–48 小时）；"
            "**B 类**中周转品走已在途库存（7–14 天）；**C 类**定制/进口品按单排柜（30–60 天）；"
            "**D 类**小件配件可由美国供应商 Dropship。四类最终均交付客户，"
            "形成「轻库存 + 在途 + 按单进口 + 配件直发」的混合模式。"
        )
    if "CEO / Duke Wang" in s:
        return (
            "CEO 下辖运营、供应链、技术、销售 BD、市场五条线。"
            "运营负责仓储 3PL 协调与客服/建材顾问；供应链负责中国采购与进口合规；"
            "技术负责产品与 AI 工程；销售 BD 覆盖 B 端 GC 与 C 端大客；"
            "市场负责内容与社群获客。"
        )
    return "详见 Markdown 源文件中的流程图章节。"


PAGE_BREAK = (
    "\n\n```{=openxml}\n"
    "<w:p><w:r><w:br w:type=\"page\"/></w:r></w:p>\n"
    "```\n\n"
)


def add_chapter_page_breaks(text: str) -> str:
    """Start each major section (##) on a new page for commercial layout."""
    return re.sub(r"^## ", PAGE_BREAK + "## ", text, flags=re.MULTILINE)


def preprocess_markdown(text: str, url_map: dict[str, str]) -> str:
    for url, r in url_map.items():
        text = text.replace(url, r)

    text = remove_milestone_section(text)

    def mermaid_replacer(match: re.Match) -> str:
        return f"\n\n{mermaid_to_prose(match.group(1))}\n\n"

    text = re.sub(r"```mermaid\n(.*?)```", mermaid_replacer, text, flags=re.DOTALL)
    text = re.sub(r"\[([^\]]+)\]\(#[^)]+\)", r"\1", text)
    text = text.replace("（文字描述）", "")
    text = add_chapter_page_breaks(text)
    return text


def build_cover_and_gallery(paths: dict[str, Path]) -> str:
    cover = f"""# HDW 建材城 · 商业策划书

**HDW LLC — 构筑理想空间，智造美好生活**

![封面 Hero]({rel(paths['hero'])})

**总部：** 5134 Biloxi Ave., North Hollywood, CA 91601  
**联系人：** Duke Wang · 323-853-3333 · dukewang@gmail.com  
**版本：** v1.0 · 2026年6月

```{{=openxml}}
<w:p><w:r><w:br w:type="page"/></w:r></w:p>
```

"""
    gallery = "\n\n" + PAGE_BREAK + "## 附录 B · 品牌与产品视觉图库\n\n"
    for i, (name, _, caption) in enumerate(IMAGE_URLS):
        if i > 0:
            gallery += PAGE_BREAK
        gallery += f"### {caption}\n\n![{caption}]({rel(paths[name])})\n\n"
    return cover, gallery


def format_docx_commercial(docx_path: Path) -> None:
    """Apply widow/orphan control, heading keep-with-next, and table row integrity."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt

    doc = Document(str(docx_path))

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    def style_name(paragraph) -> str:
        try:
            return paragraph.style.name or ""
        except Exception:
            return ""

    for paragraph in doc.paragraphs:
        pf = paragraph.paragraph_format
        pf.widow_control = True
        name = style_name(paragraph)
        if name.startswith("Heading"):
            pf.keep_with_next = True
            pf.keep_together = True
        if paragraph._element.xpath(".//w:drawing"):
            pf.keep_together = True

    for table in doc.tables:
        if len(table.rows) > 1:
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:tblHeader")) is None:
                tr_pr.append(OxmlElement("w:tblHeader"))
        for row in table.rows:
            tr_pr = row._tr.get_or_add_trPr()
            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.widow_control = True

    doc.save(str(docx_path))


def export_docx(md_path: Path, docx_path: Path) -> None:
    pypandoc.convert_file(
        str(md_path),
        "docx",
        outputfile=str(docx_path),
        extra_args=[
            "--standalone",
            "--toc",
            "--toc-depth=2",
            f"--resource-path={DOCS}",
        ],
    )
    format_docx_commercial(docx_path)


def export_pdf_via_word(docx_path: Path, pdf_path: Path) -> bool:
    tmp_pdf = Path(os.environ.get("TEMP", str(DOCS))) / "hdw_plan_export.pdf"
    try:
        import win32com.client as win32

        word = win32.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True)
        doc.ExportAsFixedFormat(str(tmp_pdf), ExportFormat=17)
        try:
            doc.Close(False)
        except Exception:
            pass
        word.Quit()
        if tmp_pdf.exists() and tmp_pdf.stat().st_size > 1000:
            shutil.copy2(tmp_pdf, pdf_path)
            return pdf_path.exists() and pdf_path.stat().st_size > 1000
    except Exception as exc:
        log(f"Word COM ExportAsFixedFormat: {exc}")
    try:
        from docx2pdf import convert

        convert(str(docx_path), str(pdf_path))
        return pdf_path.exists() and pdf_path.stat().st_size > 1000
    except Exception as exc:
        log(f"docx2pdf: {exc}")
    return False


def main() -> int:
    if not SRC_MD.exists():
        log(f"Missing source: {SRC_MD}")
        return 1

    log("Step 1: Download images...")
    paths = download_images()
    url_map = {url: rel(paths[name]) for name, url, _ in IMAGE_URLS}

    log("Step 2: Build export markdown...")
    body = preprocess_markdown(SRC_MD.read_text(encoding="utf-8"), url_map)
    cover, gallery = build_cover_and_gallery(paths)
    WORK_MD.write_text(cover + body + gallery, encoding="utf-8")

    log("Step 3: Pandoc -> DOCX...")
    out_docx = OUT_DOCX
    tmp_docx = out_docx.with_suffix(".tmp.docx")
    export_docx(WORK_MD, tmp_docx)
    try:
        shutil.move(str(tmp_docx), str(out_docx))
    except PermissionError:
        out_docx = DOCS / "HDW商业策划书_新.docx"
        shutil.move(str(tmp_docx), str(out_docx))
        log(f"DOCX locked — saved as: {out_docx} ({out_docx.stat().st_size // 1024} KB)")
    else:
        log(f"DOCX OK: {out_docx} ({out_docx.stat().st_size // 1024} KB)")

    log("Step 4: Word -> PDF...")
    out_pdf = OUT_PDF
    if export_pdf_via_word(out_docx, out_pdf):
        log(f"PDF OK: {out_pdf} ({out_pdf.stat().st_size // 1024} KB)")
    else:
        alt_pdf = DOCS / "HDW商业策划书_新.pdf"
        if export_pdf_via_word(out_docx, alt_pdf):
            log(f"PDF locked — saved as: {alt_pdf} ({alt_pdf.stat().st_size // 1024} KB)")
            out_pdf = alt_pdf
        else:
            log("PDF export failed — DOCX is ready; open in Word and Save As PDF.")
            return 2

    for copy_docx, copy_pdf in (
        (DOCS / "商业策划书_HDW建材城.docx", DOCS / "商业策划书_HDW建材城.pdf"),
    ):
        try:
            shutil.copy2(out_docx, copy_docx)
            shutil.copy2(out_pdf, copy_pdf)
            log(f"Copied -> {copy_docx.name}, {copy_pdf.name}")
        except PermissionError:
            log(f"Skip copy (file locked): {copy_docx.name} / {copy_pdf.name}")

    WORK_MD.unlink(missing_ok=True)
    log("Done.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
